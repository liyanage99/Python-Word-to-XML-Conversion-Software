import tkinter as tk
from tkinter import filedialog
from docx import Document
import xml.etree.ElementTree as ET
from xml.etree.ElementTree import SubElement
import re
import langdetect
import datefinder
import datetime
import os
from langdetect import detect
import html
import xml.etree.ElementTree as ET
from tkinter import filedialog, messagebox


#extract bold text
def extract_bold(para):
    bold_text = []
    for run in para.runs:
        if run.bold:
            bold_text.append(run.text)
    return bold_text

def convert_to_vet_tags(text):
    return f'<vet>{text}</vet>'

#extract Italic text
def extract_italic(para):
    italic_text = []
    for run in para.runs:
        if run.italic:
            italic_text.append(run.text)
    return italic_text

# Function to convert italic text into <it> tags
def convert_to_it_tags_1(text):
    return f'<it>{text}</it>'

def convert_to_it_tags_2(text):
    return f'<lat>{text}</lat>'


# extract date from title 
def find_dates(text):
    dates = list(datefinder.find_dates(text))
    return dates
 
def extract_date(title):
    match = re.search(r'\b\d{1,2}(st|nd|rd|th|er|ème)? \w+ \d{4}\b', title)
    if match:
        date_string = match.group()
        print("Found date:", date_string)
        return date_string
    else:
        print("No valid date string found in the title.")
        return None
     
def detect_language(text):
    try:
        lang = langdetect.detect(text)
        return lang
    except:
        return None


# start conversion process
def convert_to_xml(input_file, output_file, file_count):
    # Load the Word document
    doc = Document(input_file)
    
    output_folder = os.path.join(os.path.dirname(input_file), "output")
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
#root Element
    jurisr = ET.Element("jurisr")
        
  # create date variable    
    title = doc.paragraphs[1].text
    date = extract_date(title)

    # chrdatm section
    chrdatm = ET.SubElement(jurisr, "chrdatm")
    chardatum = f"{date}"
    chrdatm.text = chardatum

    

# Store paragraphs processed in the samenst section
    processed_paragraphs_samenst = set()

    # gretch and samenst section
    gretch = ET.SubElement(jurisr, "gretch" , aard="rvst")
    samenst = SubElement(gretch, "samenst")

    for para in doc.paragraphs[2:]:

        
        
        if para.text.strip():

            if ':' in para.text and 'Noot' not in para.text and 'Note' not in para.text:
                        # Process paragraphs that meet the existing condition
                        parts = para.text.split(':')
                        functie = SubElement(samenst, 'functie')
                        functie.text = parts[0].strip() + ':'
                        al = SubElement(samenst, 'al')
                        
                        al.text = parts[1].strip()
                        processed_paragraphs_samenst.add(para.text.strip())
            else:
                        # Handle paragraphs that don't meet the existing condition
                        casco_pattern = re.search(r'\b[A-Z][a-zA-Z]*\s(?:[a-zA-Z]+\s)*–\s[A-Z][a-zA-Z]*\b', para.text)
                        if casco_pattern:
                            # If the pattern is found, break out of the loop
                            break
                        else:
                            # Process the paragraph
                            al = SubElement(samenst, 'al')

                            italic_text = extract_italic(para)
                            if italic_text:
                                para_text = para.text
                                for text in italic_text:
                                    para_text = para.text.replace(text, convert_to_it_tags_1(text))
                                al.text = para_text
                            else:

                                al.text = para.text.strip()
                            processed_paragraphs_samenst.add(para.text.strip())    
            
    lang_jurisr = detect_language(para.text)
    if lang_jurisr:
        jurisr.set('xml_lang', lang_jurisr)            


    # Store paragraphs processed in the cascobol section
    

        
    # Maintain a set of processed paragraphs for <samenv> and <cascobol>
    processed_paragraphs_samenv = set()  

    processed_paragraphs_casco = set()

    # cascobol section
    if lang_jurisr == "nl":
        cascobol_nl = ET.SubElement(jurisr, "cascobol", xml_lang="nl")
        samenv_nl = ET.SubElement(jurisr, "samenv", xml_lang="nl")
        cascobol_fr = ET.SubElement(jurisr, "cascobol", xml_lang="fr")
        samenv_fr = ET.SubElement(jurisr, "samenv", xml_lang="fr")

    elif lang_jurisr == "fr":
        cascobol_fr = ET.SubElement(jurisr, "cascobol", xml_lang="fr")
        samenv_fr = ET.SubElement(jurisr, "samenv", xml_lang="fr")
        cascobol_nl = ET.SubElement(jurisr, "cascobol", xml_lang="nl")
        samenv_nl = ET.SubElement(jurisr, "samenv", xml_lang="nl")
    else:
        cascobol_nl = ET.SubElement(jurisr, "cascobol", xml_lang="nl")
        cascobol_fr = ET.SubElement(jurisr, "cascobol", xml_lang="fr")
        samenv_nl = ET.SubElement(jurisr, "samenv", xml_lang="nl")
        samenv_fr = ET.SubElement(jurisr, "samenv", xml_lang="fr")

    

    jurtekst_nl = SubElement(samenv_nl, 'jurtekst')
    jurtekst_fr = SubElement(samenv_fr, 'jurtekst')

    for para in doc.paragraphs:
        casco_pattern = re.search(r'\b[A-Z][a-zA-Z]*\s(?:[a-zA-Z]+\s)*–\s[A-Z][a-zA-Z]*\b', para.text)
        if casco_pattern and title not in para.text:
            # Determine language of casco text
            lang = detect_language(para.text)
            if lang == "nl":
                casco = SubElement(cascobol_nl, 'casco')
                cascobol_nl.set('xml_lang', 'nl')
            elif lang == "fr":
                casco = SubElement(cascobol_fr, 'casco')
                cascobol_fr.set('xml_lang', 'fr')
            else:
                # Default to NL if language detection fails or language is not NL or FR
                casco = SubElement(cascobol_nl, 'casco')
                cascobol_nl.set('lang', 'nl')

            # Process the paragraph and assign it to the appropriate casco element
            italic_text = extract_italic(para)
            if italic_text:
                processed_text = para.text
                for text in italic_text:
                    processed_text = processed_text.replace(text, convert_to_it_tags_2(text))
                casco.text = processed_text
            else:
                casco.text = para.text.strip()

    
            processed_paragraphs_casco.add(para.text.strip())

    # Iterate through paragraphs to identify and separate NL and FR samenv elements
    # Iterate through paragraphs
    for para in doc.paragraphs:
    # Extract anchor name to identify relevant paragraphs for samenv
        anchor_match = re.search(r'No\.\s\d{4}/\d{1,2}', para.text)
        if anchor_match is None and title not in para.text:
            if para.text.strip() not in processed_paragraphs_casco and para.text.strip() not in processed_paragraphs_samenst:
                # Determine language of samenv text
                lang_samenv = detect_language(para.text)
                if lang_samenv == "nl":
                    jurtekst = jurtekst_nl
                elif lang_samenv == "fr":
                    jurtekst = jurtekst_fr
                else:
                    # Default to NL if language detection fails or language is not NL or FR
                    jurtekst = jurtekst_nl or jurtekst_fr
                    

                # Split the text into paragraphs using regular expression to handle different line breaks
                paragraphs = re.split(r'\r?\n\r?\n', para.text.strip())

                # Create appropriate elements and append to the relevant samenv section
                for paragraph in paragraphs:
                    # Determine the number of paragraph breaks in the paragraph
                    paragraph_breaks = paragraph.count('\n')

                    # Create appropriate elements based on the number of paragraph breaks
                    if paragraph_breaks > 1:
                        # If there are one or more paragraph breaks, start a new <p> element
                        p = SubElement(jurtekst, 'p')
                        processed_lines = paragraph.split('\n')
                        for line in processed_lines:
                            al = SubElement(p, 'al')
                            al.text = line.strip()
                    else:
                        # If there is no paragraph break, create a new <al> element
                        al = SubElement(jurtekst, 'al')
                        al.text = paragraph.strip()

                    # Process the paragraph and assign it to the appropriate element
                    processed_text = paragraph.strip()
                    italic_text = extract_italic(para)
                    bold_text = extract_bold(para)

                    # Replace italic text with <it> tags
                    for text in italic_text:
                        processed_text = processed_text.replace(text, convert_to_it_tags_1(text))

                    # Replace bold text with <vet> tags
                    for text in bold_text:
                        processed_text = processed_text.replace(text, convert_to_vet_tags(text))

                    # Assign the processed text to the appropriate element
                    al.text = processed_text.strip()

                    # Add the processed paragraph to the set of processed paragraphs for the Samenv section
                    processed_paragraphs_samenv.add(para.text.strip())   




            

    # Uitispark Section  
    uitspraak = ET.SubElement(jurisr, "uitspraak")
    jurtekst = SubElement(uitspraak, "jurtekst")
    p = SubElement(jurtekst, "p")
    tabblok = SubElement(p, "tabblok")
    table = SubElement(tabblok, "table", frame="none")
    tgroup = SubElement(table, "tgroup", cols="2")
    colspec1 = SubElement(tgroup, "colspec", align="left", colname="col1", colnum="1", colsep="0", colwidth="5*", rowsep="0")
    colspec2 = SubElement(tgroup, "colspec", align="left", colname="col2", colnum="2", colsep="0", colwidth="5*", rowsep="0")
    tbody = SubElement(tgroup, "tbody", valign="top")
    row = SubElement(tbody, "row")
    entry1 = SubElement(row, "entry")
    al1 = SubElement(entry1, "al")
    al1.text = "Vous pouvez consulter le texte intégral via"
    al2 = SubElement(entry1, "al")
    figblok1 = SubElement(al2, "figblok")
    

    figuur1 = SubElement(figblok1, "figuur", bestand="QR-CODE", figure_name=f"2024_{file_count:03d}_FR.gif")
    entry2 = SubElement(row, "entry")
    al3 = SubElement(entry2, "al")
    al3.text = "U kan de integrale tekst raadplegen via"

    al4 = SubElement(entry2, "al")
    figblok2 = SubElement(al4, "figblok")

    figuur2 = SubElement(figblok2, "figuur", bestand="QR-CODE", figure_name=f"2024_{file_count:03d}_NL.gif")

    # Write XML to file
    
    tree = ET.ElementTree(jurisr)
    tree.write(output_file, encoding='utf-8', xml_declaration=True)
    print(f"Conversion completed for {input_file}")
    
 

# Write XML to file with line breaks
    xml_content = ET.tostring(jurisr, encoding='utf-8')
    xml_content = xml_content.decode('utf-8')
    xml_content = xml_content.replace('<jurisr', '<?xml version="1.0" encoding="utf-8"?>\n<jurisr')
    xml_content = xml_content.replace('<p><al /></p>', '')
    xml_content = xml_content.replace('><', '>\n<')
#replace all the tags with <it> and </it>    
    xml_content = xml_content.replace("&lt;it&gt;", "<it>")
    xml_content = xml_content.replace("&lt;/it&gt;", "</it>")
    xml_content = xml_content.replace("&lt;lat&gt;", "<lat>")
    xml_content = xml_content.replace("&lt;/lat&gt;", "</lat>")
#replace all the tags with <vet> and </vet>    
    xml_content = xml_content.replace("&lt;vet&gt;", "<vet>")
    xml_content = xml_content.replace("&lt;/vet&gt;", "</vet>")
#replace <al />
    xml_content = xml_content.replace("<jurtekst>", "<jurtekst>\n<p>")
    xml_content = xml_content.replace("</jurtekst>", "</p>\n</jurtekst>")
    xml_content = xml_content.replace("<al />\n<al>", "</p>\n<p><al>")
    xml_content = xml_content.replace("<al />\n", "")
    xml_content = xml_content.replace(">\n<p>\n</p>\n<", ">\n<")
    
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(xml_content)
    print(f"Conversion completed for {input_file}")



# select folder
def select_folder():
    foldername = filedialog.askdirectory()
    entry_folder.delete(0, tk.END)
    entry_folder.insert(0, foldername)

# convert files in the selected folder
def convert_folder():
    folder = entry_folder.get()
    if not folder:
        print("Please select a folder.")
        return
    try:
        file_count = int(entry_start_number.get())
    except ValueError:
        messagebox.showerror("Error", "Invalid starting number entered.")
        return
    for filename in os.listdir(folder):
        if filename.endswith(".docx"):
            input_file = os.path.join(folder, filename)
            output_folder = os.path.join(folder, "output")
            output_file = os.path.join(output_folder, filename.replace(".docx", ".xml"))
            convert_to_xml(input_file, output_file, file_count)
            file_count += 1



# Create the main window
root = tk.Tk()
root.title("Word to XML Converter")

# Folder selection
label_folder = tk.Label(root, text="Select a folder:")
label_folder.pack(pady=5)
entry_folder = tk.Entry(root, width=50)
entry_folder.pack(pady=5)
browse_button = tk.Button(root, text="Browse", command=select_folder)
browse_button.pack(pady=5)

# Starting number entry
label_start_number = tk.Label(root, text="Enter the starting number of QR Code:")
label_start_number.pack(pady=5)
entry_start_number = tk.Entry(root, width=10)
entry_start_number.pack(pady=5)

# Convert button
convert_button = tk.Button(root, text="Convert", command=convert_folder)
convert_button.pack(pady=5)

# Run the Tkinter event loop
root.mainloop()
